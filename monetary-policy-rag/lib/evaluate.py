import json
import numpy as np
from lib.retrieve import retrieve_chunks
from lib.retrieve import hybrid_rank
from docx import Document


def build_candidate_review(embedder, gold_queries, corpus_embedded, bm25, key_to_idx,k_rrf):
    output = []
    for item in gold_queries:
        query = item["query"]
        query_es=item["query_es"]
        comunicados = item["expected_comunicados"]

        n_max = len(comunicados) * 3

        #for each query, retrieve candidate chunks and add to a dictionary of candidates
        candidates = find_candidates_in_comunicado(embedder, query, query_es, comunicados, corpus_embedded, bm25, key_to_idx, n_max = n_max,k_rrf=k_rrf)

        candidate_list = []
        for c in candidates:
            candidate_list.append({
                "chunk_id": c["chunk_id"],
                "name": c["name"],
                "rrf_score": round(float(c["rrf"]), 3),
                "text_preview": c["text"][:600]  # first 300 chars, enough to judge relevance
            })

            

        #build output dict containing for each query, the difficulty what expected comunicados assigned, and each query's candidate list, then add an empty list for the subset of relevant chunk ids
        output.append({
            "query": query,
            "query_es":query_es,
            "difficulty": item["difficulty"],
            "expected_comunicados": comunicados,  # your rough comunicado guesses from before
            "candidates": candidate_list,
            "confirmed_related_chunks": []  # <-- you fill this in, per question
        })

    with open("eval/candidate_review.json", "w") as f:
        json.dump(output, f, indent=2)


    return output

def evaluate_retrieval(embedder, gold_queries, corpus_embedded, bm25, key_to_idx, threshold, k, k_rrf):
    precisions, recalls, rrs = [], [], []

    

    for query in gold_queries:
        if len(query['confirmed_related_chunks']) < 1:
            #if no gold responses because all queries are irrelevant, metrics lists should contain 0s, so that their length isn't 0 and average doesn't give divide by zero error
            precisions.append(0.0)
            recalls.append(0.0)
            rrs.append(0.0)
            continue
        
        gold = gold_ids(query)
        # print("gold", gold)
        retrieved = retrieved_ids(embedder, query["query"], query["query_es"],corpus_embedded, bm25, key_to_idx, threshold=threshold, k=k, k_rrf=k_rrf)
        # print("retrieved", retrieved)



        p = precision_at_k(retrieved, gold, k)
        rec = recall_at_k(retrieved, gold, k)
        rr = reciprocal_rank(retrieved, gold)

        precisions.append(p)
        if rec is not None:
            recalls.append(rec)
        if rr is not None:
            rrs.append(rr)

        print(f"{query['query'][:60]:60} | P@{k}={p:.2f}  R@{k}={'N/A' if rec is None else f'{rec:.2f}'}  RR={'N/A' if rr is None else f'{rr:.2f}'}")


    print(f"\nMean Precision@{k}: {sum(precisions)/len(precisions):.3f}")
    print(f"Mean Recall@{k}:    {sum(recalls)/len(recalls):.3f}")
    print(f"MRR:                {sum(rrs)/len(rrs):.3f}")

    return precisions, recalls, rrs

def find_candidates_in_comunicado(embedder, query, query_es, expected_comunicados, corpus_embedded, bm25, key_to_idx, n_max, k_rrf):
    query_vec = embedder.encode(query, normalize_embeddings=True)
    all_chunks = []
    for entry in corpus_embedded:
        if entry["name"] not in expected_comunicados:
            continue
        all_chunks.append(entry)
        entry['score'] = np.dot(query_vec, entry['embedding'])

    ranked = hybrid_rank(all_chunks, query_es, bm25, key_to_idx, k_rrf=k_rrf)

    return ranked[:n_max]

def export_candidate_review_doc(gold_queries, corpus_embedded, filepath="chunk_review.docx"):
    text_lookup = {(c['name'], c['chunk_id']): c['text'] for c in corpus_embedded}

    doc = Document()
    doc.add_heading('Candidate Chunk Review', level=1)

    for item in gold_queries:
        if not item['candidates']:
            continue  # skip irrelevant queries with no candidates

        doc.add_heading(item['query'], level=2)
        doc.add_paragraph(f"Difficulty: {item['difficulty']}")

        for c in item['candidates']:
            key = (c['name'], c['chunk_id'])
            full_text = text_lookup.get(key, "[text not found]")
            doc.add_paragraph(f"{c['name']} — chunk {c['chunk_id']} (rrf={c['rrf_score']})", style='Intense Quote')
            doc.add_paragraph(full_text)

        doc.add_page_break()

    doc.save(filepath)
    print(f"Saved {len(gold_queries)} queries to {filepath}")


def gold_ids(r):
        #given list of confirmed related chunks, return the unique identifier (comunicado + id)
    return {f"{c['name']}_{c['chunk_id']}" for c in r["confirmed_related_chunks"]}


def retrieved_ids(embedder, query, query_es, corpus_embedded, bm25, key_to_idx, threshold, k, k_rrf):
    """Run the REAL, unrestricted retrieve() and return ranked chunk ids."""
    hits = retrieve_chunks(embedder, query, query_es, corpus_embedded, bm25, key_to_idx,threshold=threshold, k=k, k_rrf=k_rrf)
    return [f"{c['name']}_{c['chunk_id']}" for c in hits]

def precision_at_k(retrieved, gold, k):
    #what share of the retrieved top k chunks are relevant 
    top_k = retrieved[:k]
    if not top_k:
        return 0.0
    hits = sum(1 for rid in top_k if rid in gold)
    return hits/len(top_k)

def recall_at_k(retrieved, gold, k):
    #what share of the confirmed relevant chunks were retrieved?  
    if not gold:
        return None 
    top_k = retrieved[:k]
    hits = sum(1 for rid in top_k if rid in gold)

    return hits/len(gold)

def reciprocal_rank(retrieved, gold):
    if not gold:
        return None
    for rank, rid in enumerate(retrieved, start=1):
        if rid in gold:
            return 1.0/rank
    return 0.0