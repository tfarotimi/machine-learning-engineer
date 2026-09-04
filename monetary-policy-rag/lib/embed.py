from docx import Document
from sentence_transformers import SentenceTransformer

#Embed each chunk in the vector space
def embed_chunks(embedder, chunks):
    to_embed = []
    embedded = []

    for comunicado in chunks:
        for i in chunks[comunicado]:
            #create a list of dicts (containing each chunk and its embedding)
            embedded.append({"name":comunicado,
                        "chunk_id":i,
                        "text": chunks[comunicado][i]})
            
            #append each chunk to a flat list for encodding
            to_embed.append(chunks[comunicado][i])

    embeddings = embedder.encode(to_embed, normalize_embeddings=True) # normalized to unit length so retrieval reduces to a plain dot product (cosine theta = a.b when magnitude = 1)


    for i, vec  in enumerate(embeddings):
            embedded[i]["embedding"] = vec


    return embedded

#load embedder model
def load_embedder(model_name='sentence-transformers/paraphrase-multilingual-mpnet-base-v2'):
    return SentenceTransformer(model_name)


def export_full_corpus_doc(corpus_embedded, filepath="full_corpus.docx"):
    doc = Document()
    doc.add_heading('Full Corpus — Chunked Text by comunicado', level=1)

    # Group chunks by comunicado, preserving chunk order within each comunicado
    comunicados = {}
    for c in corpus_embedded:
        comunicados.setdefault(c['name'], []).append(c)

    for comunicado_name in sorted(comunicados.keys()):
        chunks = sorted(comunicados[comunicado_name], key=lambda c: c['chunk_id'])
        doc.add_heading(comunicado_name, level=2)
        for c in chunks:
            doc.add_paragraph(f"Chunk {c['chunk_id']}", style='Intense Quote')
            doc.add_paragraph(c['text'])
        doc.add_page_break()

    doc.save(filepath)
    print(f"Saved {len(comunicados)} comunicados, {len(corpus_embedded)} chunks to {filepath}")